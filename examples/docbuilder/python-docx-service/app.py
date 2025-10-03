from flask import Flask, request, jsonify, send_file
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_UNDERLINE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import os, uuid, io, base64

app = Flask(__name__)
OUT_DIR = '/tmp/docx_outputs'
os.makedirs(OUT_DIR, exist_ok=True)

def set_font(run, name='Calibri', size=11, bold=False, italic=False, underline=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if underline:
        run.underline = True
    # Fix font for East Asian languages
    rPr = run._element.rPr
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), name)
    rPr.append(rFonts)

def add_cover_page(doc, title, author, logo_path=None):
    # Add logo if provided
    if logo_path and os.path.exists(logo_path):
        doc.add_picture(logo_path, width=Inches(2))
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run(title)
    set_font(run, size=28, bold=True)
    doc.add_paragraph()
    p2 = doc.add_paragraph()
    p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run2 = p2.add_run(f"Author: {author}")
    set_font(run2, size=14, italic=True)
    p3 = doc.add_paragraph()
    p3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run3 = p3.add_run(f"Date: {datetime.now().strftime('%B %d, %Y')}")
    set_font(run3, size=12)
    doc.add_page_break()

def add_header_footer(doc, title):
    section = doc.sections[0]

    # Header with title aligned left
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.text = title
    paragraph.style = 'Header'

    # Footer with page number centered
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # Add page number field
    run = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def add_table_of_contents(doc):
    # Insert TOC field for Word to auto-generate TOC
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = "TOC \\o \"1-3\" \\h \\z \\u"  # TOC field code
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:t')
    fldChar3.text = "Right-click to update field."
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')
    r_element = run._r
    r_element.append(fldChar1)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar3)
    r_element.append(fldChar4)

def add_section(doc, title, content, level=1):
    heading = doc.add_heading(title, level=level)
    # Customize heading font style
    run = heading.runs[0]
    set_font(run, size=16 - (level * 2), bold=True)

    for paragraph in content.split('\n\n'):
        p = doc.add_paragraph(paragraph)
        set_font(p.runs[0] if p.runs else p.add_run(), size=11)

def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
        set_font(p.runs[0], size=11)

def add_numbered_list(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style='List Number')
        set_font(p.runs[0], size=11)

def add_blockquote(doc, text):
    p = doc.add_paragraph(text, style='Intense Quote')
    set_font(p.runs[0], italic=True, size=11)

def add_footnote(paragraph, text):
    # python-docx doesn't support footnotes natively,
    # so we simulate by appending superscript number with reference text at bottom
    # This is a hack, for full support you might use 'python-docx' extensions or direct XML manipulation
    run = paragraph.add_run(' [1]')
    run.font.superscript = True
    # Append footnote text at the end of document
    doc = paragraph.part.document
    doc.add_paragraph('1. ' + text, style='Footnote Text')

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
    for row in rows:
        row_cells = table.add_row().cells
        for i, cell in enumerate(row):
            row_cells[i].text = str(cell)
    # Style the table
    table.style = 'Light List Accent 1'

def add_diagram_image(doc, image_bytes):
    stream = io.BytesIO(image_bytes)
    doc.add_picture(stream, width=Inches(4))

@app.route('/generate', methods=['POST'])
def generate():
    payload = request.get_json()
    title = payload.get('title', 'Document Title')
    author = payload.get('author', 'Author')
    outline = payload.get('outline', {})
    content = payload.get('content', {})
    diagrams = payload.get('diagrams', {})
    bullets = payload.get('bullets', {})
    numbered = payload.get('numbered', {})
    tables = payload.get('tables', {})

    doc = Document()

    add_cover_page(doc, title, author, logo_path=None)
    add_header_footer(doc, title)
    add_table_of_contents(doc)
    doc.add_page_break()

    # Add main content sections
    for sec in outline.get('sections', []):
        sec_text = content.get(sec, '')
        add_section(doc, sec, sec_text, level=1)

        if sec in bullets:
            add_bullets(doc, bullets[sec])
        if sec in numbered:
            add_numbered_list(doc, numbered[sec])
        if sec in tables:
            table_data = tables[sec]
            add_table(doc, table_data.get('headers', []), table_data.get('rows', []))
        if sec in diagrams:
            doc.add_heading('Diagram', level=2)
            # For demonstration, assume diagrams[sec] is base64 PNG data
            try:
                img_bytes = base64.b64decode(diagrams[sec])
                add_diagram_image(doc, img_bytes)
            except Exception as e:
                doc.add_paragraph(f"Could not render diagram: {str(e)}")

    doc.add_page_break()
    doc.add_heading('Appendix', level=1)
    doc.add_paragraph('Generated with Document Builder (demo).', style='Intense Quote')

    filename = f"{uuid.uuid4().hex}.docx"
    outpath = os.path.join(OUT_DIR, filename)
    doc.save(outpath)
    return jsonify({"path": outpath})

@app.route('/download')
def download():
    path = request.args.get('path')
    if not path or not os.path.exists(path):
        return 'not found', 404
    return send_file(path, as_attachment=True)

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000)
